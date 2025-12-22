"""
DubSync DOCX Export Plugin

Word document export plugin for professional dubbing scripts.
"""

from pathlib import Path
from typing import List, Dict, Any, Optional
from datetime import datetime

from PySide6.QtWidgets import (
    QWidget, QVBoxLayout, QHBoxLayout, QFormLayout,
    QLabel, QPushButton, QCheckBox, QComboBox, QLineEdit,
    QGroupBox, QFileDialog, QMessageBox, QSpinBox
)
from PySide6.QtCore import Qt
from PySide6.QtGui import QAction

from dubsync.plugins.base import ExportPlugin, UIPlugin, PluginInfo, PluginType, PluginDependency
from dubsync.models.project import Project
from dubsync.models.cue import Cue
from dubsync.utils.time_utils import ms_to_timecode
from dubsync.utils.constants import CueStatus
from dubsync.i18n import t

from typing import TYPE_CHECKING
if TYPE_CHECKING:
    from dubsync.ui.main_window import MainWindow


class DOCXExportOptionsWidget(QWidget):
    """DOCX export options widget."""
    
    def __init__(self, plugin: 'DOCXExportPlugin', parent=None):
        super().__init__(parent)
        self.plugin = plugin
        self._setup_ui()
    
    def _setup_ui(self):
        layout = QVBoxLayout(self)
        layout.setContentsMargins(8, 8, 8, 8)
        layout.setSpacing(8)

        # Header
        header = QLabel(t("plugins.docx_export.header"))
        header.setStyleSheet("font-size: 14px; font-weight: bold;")
        layout.addWidget(header)

        # Content settings
        content_group = QGroupBox(t("plugins.docx_export.content_group"))
        content_form = QFormLayout(content_group)

        self.include_timecodes_cb = self._extracted_from__setup_ui_15(
            True, content_form, "plugins.docx_export.include_timecodes"
        )
        self.include_source_cb = self._extracted_from__setup_ui_15(
            True, content_form, "plugins.docx_export.include_source"
        )
        self.include_character_cb = self._extracted_from__setup_ui_15(
            True, content_form, "plugins.docx_export.include_character"
        )
        self.include_notes_cb = self._extracted_from__setup_ui_15(
            True, content_form, "plugins.docx_export.include_notes"
        )
        self.include_sfx_cb = self._extracted_from__setup_ui_15(
            False, content_form, "plugins.docx_export.include_sfx"
        )
        self.include_status_cb = self._extracted_from__setup_ui_15(
            True, content_form, "plugins.docx_export.include_status"
        )
        layout.addWidget(content_group)

        # Formatting settings
        format_group = QGroupBox(t("plugins.docx_export.format_group"))
        format_form = QFormLayout(format_group)

        self.style_combo = QComboBox()
        self.style_combo.addItem(t("plugins.docx_export.style_table"), "table")
        self.style_combo.addItem(t("plugins.docx_export.style_script"), "script")
        format_form.addRow(t("plugins.docx_export.style"), self.style_combo)

        self.font_size_spin = QSpinBox()
        self.font_size_spin.setRange(8, 16)
        self.font_size_spin.setValue(11)
        self.font_size_spin.setSuffix(" pt")
        format_form.addRow(t("plugins.docx_export.font_size"), self.font_size_spin)

        self.page_break_cb = self._extracted_from__setup_ui_15(
            False, format_form, "plugins.docx_export.page_break"
        )
        layout.addWidget(format_group)

        # Export button
        self.export_btn = QPushButton(t("plugins.docx_export.export_button"))
        self.export_btn.clicked.connect(self._on_export)
        layout.addWidget(self.export_btn)

        layout.addStretch()

    # TODO Rename this here and in `_setup_ui`
    def _extracted_from__setup_ui_15(self, arg0, arg1, arg2):
        result = QCheckBox()
        result.setChecked(arg0)
        arg1.addRow(t(arg2), result)

        return result

    def get_options(self) -> Dict[str, Any]:
        """Get export options."""
        return {
            "include_timecodes": self.include_timecodes_cb.isChecked(),
            "include_source": self.include_source_cb.isChecked(),
            "include_character": self.include_character_cb.isChecked(),
            "include_notes": self.include_notes_cb.isChecked(),
            "include_sfx": self.include_sfx_cb.isChecked(),
            "include_status": self.include_status_cb.isChecked(),
            "style": self.style_combo.currentData(),
            "font_size": self.font_size_spin.value(),
            "page_break": self.page_break_cb.isChecked(),
        }
    
    def _on_export(self):
        """Export button clicked."""
        if not self.plugin._main_window:
            return

        pm = getattr(self.plugin._main_window, 'project_manager', None)
        if not pm or not pm.is_open or not pm.project:
            QMessageBox.warning(
                self,
                t("plugins.docx_export.error_title"),
                t("plugins.docx_export.no_project")
            )
            return

        # Get save path
        default_name = pm.project.name or "export"
        default_path = f"{default_name}.docx"
        
        file_path, _ = QFileDialog.getSaveFileName(
            self,
            t("plugins.docx_export.save_title"),
            default_path,
            t("plugins.docx_export.file_filter")
        )

        if not file_path:
            return

        # Perform export
        try:
            cues = pm.get_cues()
            options = self.get_options()
            self.plugin.export(pm.project, cues, file_path, **options)
            
            QMessageBox.information(
                self,
                t("plugins.docx_export.success_title"),
                t("plugins.docx_export.success_message", path=file_path)
            )
        except ImportError as e:
            QMessageBox.critical(
                self,
                t("plugins.docx_export.error_title"),
                t("plugins.docx_export.missing_dependency")
            )
        except Exception as e:
            QMessageBox.critical(
                self,
                t("plugins.docx_export.error_title"),
                t("plugins.docx_export.error_message", error=str(e))
            )


class DOCXExportPlugin(ExportPlugin, UIPlugin):
    """
    DOCX Export Plugin.
    
    Exports project to a Word document with professional formatting.
    Supports two styles: table-based and script-based layout.
    """

    def __init__(self):
        super().__init__()
        self._widget: Optional[DOCXExportOptionsWidget] = None

    @property
    def info(self) -> PluginInfo:
        return PluginInfo(
            id="docx_export",
            name=t("plugins.docx_export.name"),
            version="1.0.0",
            author="Levente Kulacsy",
            description=t("plugins.docx_export.description"),
            plugin_type=PluginType.EXPORT,
            dependencies=[
                PluginDependency(
                    package_name="python-docx",
                    min_version="0.8.0"
                )
            ],
            icon="📝",
            readme_path="README.md"
        )
    
    def get_ui_widget(self) -> Optional[QWidget]:
        """Get UI widget for dock."""
        if not self._widget:
            self._widget = DOCXExportOptionsWidget(self)
        return self._widget
    
    def export(self, project: Project, cues: List[Cue], output_path: str, **options) -> bool:
        """
        Export project to DOCX.
        
        Args:
            project: Project object
            cues: List of cues
            output_path: Output file path
            **options: Export options
            
        Returns:
            True if successful
        """
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.table import WD_TABLE_ALIGNMENT
        except ImportError as e:
            raise ImportError("python-docx library is required for DOCX export") from e

        style = options.get("style", "table")
        font_size = options.get("font_size", 11)

        # Create document
        doc = Document()

        # Set default style
        style_obj = doc.styles['Normal']
        font = style_obj.font
        font.name = 'Calibri'
        font.size = Pt(font_size)

        # Add title
        title = doc.add_heading(project.title or t("plugins.docx_export.untitled"), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add metadata
        self._add_metadata(doc, project, options)

        # Add content based on style
        if style == "table":
            self._export_table_style(doc, cues, options)
        else:
            self._export_script_style(doc, cues, options)

        # Save document
        doc.save(output_path)

        return True
    
    def _add_metadata(self, doc, project: Project, options: Dict[str, Any]) -> None:
        """Add project metadata to document."""
        from docx.shared import Pt
        
        meta_items = []
        
        if project.series_title:
            meta_items.append(f"{t('dialogs.project_settings.series')}: {project.series_title}")
        
        if project.season or project.episode:
            season_ep = []
            if project.season:
                season_ep.append(f"S{project.season}")
            if project.episode:
                season_ep.append(f"E{project.episode}")
            if project.episode_title:
                season_ep.append(f"- {project.episode_title}")
            meta_items.append(" ".join(season_ep))
        
        if project.translator:
            meta_items.append(f"{t('dialogs.project_settings.translator')}: {project.translator}")
        
        if project.editor:
            meta_items.append(f"{t('dialogs.project_settings.editor')}: {project.editor}")
        
        # Add export date
        meta_items.append(f"{t('plugins.docx_export.export_date')}: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        
        if meta_items:
            for item in meta_items:
                p = doc.add_paragraph(item)
                p.runs[0].font.size = Pt(10)
                p.runs[0].font.italic = True
            doc.add_paragraph()  # Spacer
    
    def _export_table_style(self, doc, cues: List[Cue], options: Dict[str, Any]) -> None:
        """Export cues as a table."""
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import nsdecls
        from docx.oxml import parse_xml
        
        include_timecodes = options.get("include_timecodes", True)
        include_source = options.get("include_source", True)
        include_character = options.get("include_character", True)
        include_notes = options.get("include_notes", True)
        include_sfx = options.get("include_sfx", False)
        include_status = options.get("include_status", True)
        
        # Build column headers
        headers = ["#"]
        if include_timecodes:
            headers.extend([t("cue_list.columns.time_in"), t("cue_list.columns.time_out")])
        if include_character:
            headers.append(t("cue_list.columns.character"))
        if include_source:
            headers.append(t("plugins.docx_export.source_column"))
        headers.append(t("plugins.docx_export.translation_column"))
        if include_notes:
            headers.append(t("plugins.docx_export.notes_column"))
        if include_sfx:
            headers.append(t("cue_editor.sfx"))
        if include_status:
            headers.append(t("cue_list.columns.status"))
        
        # Create table
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        
        # Header row
        hdr_cells = table.rows[0].cells
        for idx, header in enumerate(headers):
            hdr_cells[idx].text = header
            hdr_cells[idx].paragraphs[0].runs[0].bold = True
            
            # Set header background color
            shading = parse_xml(
                f'<w:shd {nsdecls("w")} w:fill="4472C4"/>'
            )
            hdr_cells[idx]._tc.get_or_add_tcPr().append(shading)
            hdr_cells[idx].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        
        # Status colors
        status_colors = {
            CueStatus.NEW: "E0E0E0",
            CueStatus.TRANSLATED: "81C784",
            CueStatus.NEEDS_REVISION: "FFB74D",
            CueStatus.APPROVED: "4CAF50",
        }
        
        # Data rows
        for cue in cues:
            row_cells = table.add_row().cells
            col_idx = 0
            
            # Index
            row_cells[col_idx].text = str(cue.cue_index)
            col_idx += 1
            
            if include_timecodes:
                row_cells[col_idx].text = ms_to_timecode(cue.time_in_ms)[:8]
                col_idx += 1
                row_cells[col_idx].text = ms_to_timecode(cue.time_out_ms)[:8]
                col_idx += 1
            
            if include_character:
                row_cells[col_idx].text = cue.character_name or "-"
                col_idx += 1
            
            if include_source:
                row_cells[col_idx].text = cue.source_text
                col_idx += 1
            
            # Translation
            row_cells[col_idx].text = cue.translated_text or ""
            col_idx += 1
            
            if include_notes:
                row_cells[col_idx].text = cue.notes or ""
                col_idx += 1
            
            if include_sfx:
                row_cells[col_idx].text = cue.sfx_notes or ""
                col_idx += 1
            
            if include_status:
                status_text = t(f"status.{cue.status.name.lower()}")
                row_cells[col_idx].text = status_text
                
                # Color the status cell
                color = status_colors.get(cue.status, "E0E0E0")
                shading = parse_xml(
                    f'<w:shd {nsdecls("w")} w:fill="{color}"/>'
                )
                row_cells[col_idx]._tc.get_or_add_tcPr().append(shading)
    
    def _export_script_style(self, doc, cues: List[Cue], options: Dict[str, Any]) -> None:
        """Export cues in screenplay format."""
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        include_timecodes = options.get("include_timecodes", True)
        include_source = options.get("include_source", True)
        include_notes = options.get("include_notes", True)
        page_break = options.get("page_break", False)
        
        for idx, cue in enumerate(cues):
            # Character name (centered, uppercase)
            if cue.character_name:
                char_para = doc.add_paragraph()
                char_run = char_para.add_run(cue.character_name.upper())
                char_run.bold = True
                char_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Timecode (small, gray)
            if include_timecodes:
                time_para = doc.add_paragraph()
                time_text = f"[{ms_to_timecode(cue.time_in_ms)[:8]} - {ms_to_timecode(cue.time_out_ms)[:8]}]"
                time_run = time_para.add_run(time_text)
                time_run.font.size = Pt(9)
                time_run.font.color.rgb = RGBColor(128, 128, 128)
                time_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Source text (italic, if included)
            if include_source and cue.source_text:
                source_para = doc.add_paragraph()
                source_run = source_para.add_run(cue.source_text)
                source_run.italic = True
                source_run.font.color.rgb = RGBColor(100, 100, 100)
            
            # Translation text
            if cue.translated_text:
                trans_para = doc.add_paragraph(cue.translated_text)
                trans_para.paragraph_format.left_indent = Pt(36)
                trans_para.paragraph_format.right_indent = Pt(36)
            
            # Notes (if included)
            if include_notes and cue.notes:
                notes_para = doc.add_paragraph()
                notes_run = notes_para.add_run(f"[{t('cue_editor.notes')} {cue.notes}]")
                notes_run.font.size = Pt(9)
                notes_run.font.italic = True
                notes_run.font.color.rgb = RGBColor(0, 100, 200)
            
            # Add separator or page break
            if idx < len(cues) - 1:
                if page_break:
                    doc.add_page_break()
                else:
                    doc.add_paragraph("─" * 40).alignment = WD_ALIGN_PARAGRAPH.CENTER
                    doc.add_paragraph()  # Spacer
    
    def get_export_formats(self) -> List[str]:
        """Get supported export formats."""
        return ["docx"]
    
    def get_menu_action(self) -> Optional[QAction]:
        """Get menu action."""
        action = QAction(t("plugins.docx_export.menu_action"), None)
        action.triggered.connect(self._on_menu_export)
        return action
    
    def _on_menu_export(self):
        """Menu export triggered."""
        if self._widget:
            self._widget._on_export()


# Plugin instance for registration
plugin = DOCXExportPlugin()
